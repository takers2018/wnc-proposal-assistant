import io, re
from docx import Document

# very light Markdown→Docx pass:
# - #/##/### → heading 1/2/3
# - lines starting with "- " → bullet list
# - everything else → paragraph
# - **bold** and *italics* in-line (best-effort)

_bold_pat   = re.compile(r"\*\*(.+?)\*\*")
_italic_pat = re.compile(r"\*(.+?)\*")

def _add_markdown_line(doc: Document, line: str):
    if not line.strip():
        doc.add_paragraph("")
        return

    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
        return
    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
        return
    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
        return

    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        text = line[2:].strip()
    else:
        p = doc.add_paragraph()
        text = line

    # simple inline formatting (**bold**, *italics*)
    # split on bold first, then italics inside each segment
    def emit_inline(dst_paragraph, txt):
        pos = 0
        for m in _bold_pat.finditer(txt):
            if m.start() > pos:
                dst_paragraph.add_run(txt[pos:m.start()])
            bold_seg = m.group(1)
            # handle italics inside bold segment
            ipos = 0
            for im in _italic_pat.finditer(bold_seg):
                if im.start() > ipos:
                    r = dst_paragraph.add_run(bold_seg[ipos:im.start()]); r.bold = True
                ir = dst_paragraph.add_run(im.group(1)); ir.bold = True; ir.italic = True
                ipos = im.end()
            if ipos < len(bold_seg):
                r = dst_paragraph.add_run(bold_seg[ipos:]); r.bold = True
            pos = m.end()
        if pos < len(txt):
            # remaining, allow *italics*
            t = txt[pos:]
            ipos = 0
            for im in _italic_pat.finditer(t):
                if im.start() > ipos:
                    dst_paragraph.add_run(t[ipos:im.start()])
                ir = dst_paragraph.add_run(im.group(1)); ir.italic = True
                ipos = im.end()
            if ipos < len(t):
                dst_paragraph.add_run(t[ipos:])

    emit_inline(p, text)

def markdownish_to_docx_bytes(title: str, markdown_text: str) -> bytes:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    # split paragraphs but keep blank lines
    for raw_line in markdown_text.replace("\r\n","\n").split("\n"):
        _add_markdown_line(doc, raw_line.rstrip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
